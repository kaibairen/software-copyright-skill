#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成传统康普茶数字化生产管理平台V1.0软件使用手册
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def setup_header_with_page_number(section, software_name, version):
    """设置页眉：软件名称+版本号居中，页码居中，带下边框线"""
    header = section.header
    for para in header.paragraphs:
        para.clear()

    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.tab_stops.add_tab_stop(Inches(3.0), WD_ALIGN_PARAGRAPH.CENTER)

    run_text = para.add_run(f'\t{software_name} {version}\t')
    run_text.font.size = Pt(9)
    run_text.font.name = '宋体'
    run_text._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    run_page = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)

    run_page.font.size = Pt(9)
    run_page.font.name = '宋体'
    run_page._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')

    pBdr.append(bottom)
    pPr.append(pBdr)

def set_chinese_font(run, font_name='宋体', font_size=None, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True

def add_cover_page(doc, software_name, version):
    """添加封面页"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(software_name)
    set_chinese_font(run, '宋体', 22, True)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_para.add_run(version)
    set_chinese_font(run, '宋体', 18, True)

def parse_heading_level(line):
    """解析标题级别"""
    if line.startswith('#### '):
        text = line[4:].strip().replace('**', '')
        return (4, text)
    if line.startswith('### '):
        text = line[3:].strip().replace('**', '')
        return (3, text)
    if line.startswith('## '):
        text = line[2:].strip().replace('**', '')
        if re.match(r'[一二三]、', text):
            return (1, text)
        return (2, text)
    if re.match(r'^\*\*[一二三]、', line):
        return (1, line.replace('**', '').strip())
    if re.match(r'^\*\*3\.\d+\s+', line):
        return (2, line.replace('**', '').strip())
    if re.match(r'^\*\*3\.\d+\.\d+\s+', line):
        return (3, line.replace('**', '').strip())
    if re.match(r'^\*\*2\.\d+\s+', line):
        return (2, line.replace('**', '').strip())
    if line.startswith('**') and line.endswith('**'):
        return (4, line.replace('**', '').strip())
    return (0, line)

def get_heading_font_size(level):
    """根据标题级别返回字体大小"""
    sizes = {1: 18, 2: 16, 3: 15, 4: 14}
    return sizes.get(level, 12)

def add_heading(doc, text, level):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)

    font_size = get_heading_font_size(level)
    set_chinese_font(run, '宋体', font_size, True)
    p.paragraph_format.space_after = Pt(6)

def add_body_paragraph(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    clean_text = text.replace('**', '').replace('*', '')

    run = p.add_run(clean_text)
    set_chinese_font(run, '宋体', 12, False)

    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.5

def add_image_placeholder(doc, image_num):
    """添加图片占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[图{image_num} 功能截图位置]')
    set_chinese_font(run, '宋体', 10, False)
    run.font.italic = True

def generate_user_manual():
    """生成用户手册主函数"""
    print("="*60)
    print("传统康普茶数字化生产管理平台V1.0 - 使用手册生成器")
    print("="*60)

    software_name = '传统康普茶数字化生产管理平台'
    version = 'V1.0'
    markdown_path = 'manual/manual.md'
    output_path = 'materials/传统康普茶数字化生产管理平台V1.0使用手册.docx'

    doc = Document()

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    section.header_distance = Inches(0.24)
    section.footer_distance = Inches(0.69)

    setup_header_with_page_number(section, software_name, version)

    add_cover_page(doc, software_name, version)

    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    image_count = 0

    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        if line.startswith('# '):
            continue

        if '[IMAGE_MASK]' in line:
            image_count += 1
            add_image_placeholder(doc, image_count)
            continue

        level, text = parse_heading_level(line)

        if level > 0:
            add_heading(doc, text, level)
        else:
            add_body_paragraph(doc, text)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"\n[SUCCESS] 使用手册已生成!")
    print(f"[FILE] 保存路径: {output_path}")
    print(f"[NOTICE] 文档中包含 {image_count} 个图片占位符")
    print(f"[REMINDER] 请对 static-example.html 进行截图并粘贴到docx文档中")
    print("="*60)

if __name__ == '__main__':
    generate_user_manual()
