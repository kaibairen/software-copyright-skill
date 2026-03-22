#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
填写软著申请表
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

def set_cell_font(cell):
    """设置单元格内容的字体为宋体8号"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(8)

def fill_application_form():
    """填写申请表主函数"""
    print("="*60)
    print("填写软著申请表")
    print("="*60)

    project_info = {
        'software_name': '传统康普茶数字化生产管理平台',
        'version': 'V1.0',
        'short_name': '康普茶生产管理系统',
        'software_type': '应用软件',
        'completion_date': '2025年01月21日',
        'dev_hardware': 'CPU: 双核2.0GHz及以上, 内存: 8GB及以上, 硬盘: 100GB及以上',
        'runtime_hardware': 'CPU: 双核2.0GHz及以上, 内存: 4GB及以上, 硬盘: 50GB及以上',
        'dev_os': 'Windows 10/11, Linux (Ubuntu 20.04+), macOS',
        'dev_tools': 'Python 3.8+, Flask 2.3+, Vue.js 3.x, Element Plus, MySQL, InfluxDB',
        'runtime_platform': 'Windows 10/11, Linux, macOS, Web浏览器(Chrome, Firefox, Edge)',
        'runtime_support': 'Python运行环境, MySQL数据库, Web浏览器',
        'programming_language': 'Python(后端核心), HTML5(前端), CSS3(样式), JavaScript ES6+(交互), SQL(数据库), Jinja2(模板), MQTT(IoT通信)',
        'code_lines': '约3500行(前端约1800行,后端约1700行)',
        'dev_purpose': '通过数字化手段实现康普茶生产全流程的信息化管理，包括原料管理、发酵过程监控、批次追踪、库存管理等功能，提高生产效率、保证产品质量、降低生产成本',
        'application_field': '食品加工、饮料制造、发酵工程、生产管理',
        'main_functions': '原料信息管理模块实现原料档案建立、分类管理、库存监控等功能；发酵流程监控模块支持第一发酵和第二发酵全过程数据录入和实时监控；批次追踪系统模块实现从原料到成品的完整追溯链路，支持向前溯源和向后追踪；库存预警管理模块提供智能库存预警和补货建议；实时数据监控模块通过IoT技术实现发酵参数实时采集和可视化展示；智能工艺优化模块基于历史数据分析提供参数优化建议；任务提醒功能模块自动生成任务提醒并推送给相关人员；用户权限管理模块实现基于角色的访问控制。',
        'technical_features': '本软件采用数据库驱动架构设计，通过SQLAlchemy ORM框架实现数据持久化，包含完整的数据库模型定义和数据初始化服务。基于规则的预警系统从数据库动态读取预警规则进行实时监控。采用MQTT协议实现IoT设备数据采集，InfluxDB时序数据库存储传感器数据。采用JWT Token实现身份认证，使用bcrypt加密密码确保数据安全。采用RESTful API架构实现前后端分离，响应式Web设计支持多终端访问。实现加权评分模型进行发酵成熟度计算。系统包含完整的CRUD操作和数据验证机制，支持WebSocket实时数据推送。'
    }

    template_path = '../.claude/skills/software-copyright-design/asserts/申请表模板.docx'
    output_path = 'materials/传统康普茶数字化生产管理平台V1.0申请表.docx'

    print("\n[INFO] 项目信息:")
    for key, value in project_info.items():
        print(f"  - {key}: {value}")

    if not os.path.exists(template_path):
        print(f"\n[WARNING] 申请表模板文件不存在: {template_path}")
        print("[INFO] 将生成申请表信息文档，您可以手动填入官方申请表")

        # 生成信息文档
        info_path = 'materials/申请表填写信息.txt'
        with open(info_path, 'w', encoding='utf-8') as f:
            f.write("="*60 + "\n")
            f.write("传统康普茶数字化生产管理平台V1.0 - 申请表填写信息\n")
            f.write("="*60 + "\n\n")

            for key, value in project_info.items():
                f.write(f"{key}:\n")
                f.write(f"  {value}\n\n")

        print(f"\n[SUCCESS] 申请表信息已保存到: {info_path}")
        print("[提示] 请根据上述信息手动填写官方申请表")
        return

    try:
        doc = Document(template_path)
        table = doc.tables[0]

        for row in table.rows:
            for cell in row.cells:
                set_cell_font(cell)

        fill_data = {
            1: project_info['software_name'],
            2: project_info['version'],
            3: project_info['short_name'],
            4: project_info['software_type'],
            5: project_info['completion_date'],
            11: project_info['dev_hardware'],
            12: project_info['runtime_hardware'],
            13: project_info['dev_os'],
            14: project_info['dev_tools'],
            15: project_info['runtime_platform'],
            16: project_info['runtime_support'],
            17: project_info['programming_language'],
            18: project_info['code_lines'],
            19: project_info['dev_purpose'],
            20: project_info['application_field'],
            21: project_info['main_functions'],
            22: project_info['technical_features']
        }

        filled_count = 0
        for row_idx, text in fill_data.items():
            if row_idx < len(table.rows):
                cell = table.rows[row_idx].cells[1]
                current_value = cell.text.strip()

                if not current_value or current_value == '':
                    cell.text = str(text)
                    set_cell_font(cell)
                    filled_count += 1
                    print(f"[NEW] 已填写行{row_idx}: {text[:30]}...")
                else:
                    print(f"[KEEP] 保留行{row_idx}")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        print(f"\n[SUCCESS] 申请表已生成!")
        print(f"[FILE] 保存路径: {output_path}")
        print(f"[STATS] 新填写条目: {filled_count}个")
        print("="*60)

    except Exception as e:
        print(f"\n[ERROR] 生成失败: {str(e)}")
        print("[提示] 请检查申请表模板文件是否正确")

if __name__ == '__main__':
    fill_application_form()
