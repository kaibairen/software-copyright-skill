#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成软著申请表docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import os

def set_cell_font(cell, font_size=8, bold=False):
    """设置单元格内容的字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True

def generate_application_form_docx():
    """生成申请表docx主函数"""
    print("="*60)
    print("生成软著申请表docx")
    print("="*60)

    output_path = 'materials/传统康普茶数字化生产管理平台V1.0申请表.docx'

    # 创建文档
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # 添加标题
    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run('计算机软件著作权登记申请表')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)
    run.font.bold = True
    title.paragraph_format.space_after = Pt(12)

    # 创建表格
    table_data = [
        ("软件全称", "传统康普茶数字化生产管理平台"),
        ("版本号", "V1.0"),
        ("软件简称", "康普茶生产管理系统"),
        ("软件分类", "应用软件"),
        ("主要功能", "原料信息管理模块实现原料档案建立、分类管理、库存监控等功能；发酵流程监控模块支持第一发酵和第二发酵全过程数据录入和实时监控；批次追踪系统模块实现从原料到成品的完整追溯链路，支持向前溯源和向后追踪；库存预警管理模块提供智能库存预警和补货建议；实时数据监控模块通过IoT技术实现发酵参数实时采集和可视化展示；智能工艺优化模块基于历史数据分析提供参数优化建议；任务提醒功能模块自动生成任务提醒并推送给相关人员；用户权限管理模块实现基于角色的访问控制。"),
        ("技术特点", "本软件采用数据库驱动架构设计，通过SQLAlchemy ORM框架实现数据持久化，包含完整的数据库模型定义和数据初始化服务。基于规则的预警系统从数据库动态读取预警规则进行实时监控。采用MQTT协议实现IoT设备数据采集，InfluxDB时序数据库存储传感器数据。采用JWT Token实现身份认证，使用bcrypt加密密码确保数据安全。采用RESTful API架构实现前后端分离，响应式Web设计支持多终端访问。实现加权评分模型进行发酵成熟度计算。系统包含完整的CRUD操作和数据验证机制，支持WebSocket实时数据推送。"),
        ("开发完成日期", "2025年01月21日"),
        ("首次发表日期", "未发表"),
        ("开发方式", "原创"),
        ("硬件环境", "开发硬件环境：CPU: 双核2.0GHz及以上, 内存: 8GB及以上, 硬盘: 100GB及以上\n运行硬件环境：CPU: 双核2.0GHz及以上, 内存: 4GB及以上, 硬盘: 50GB及以上"),
        ("软件环境", "开发操作系统：Windows 10/11, Linux (Ubuntu 20.04+), macOS\n开发工具：Python 3.8+, Flask 2.3+, Vue.js 3.x, Element Plus, MySQL, InfluxDB\n运行平台：Windows 10/11, Linux, macOS, Web浏览器(Chrome, Firefox, Edge)\n支撑环境：Python运行环境, MySQL数据库, Web浏览器"),
        ("编程语言", "Python(后端核心), HTML5(前端), CSS3(样式), JavaScript ES6+(交互), SQL(数据库), Jinja2(模板), MQTT(IoT通信)"),
        ("源程序量", "约3500行(前端约1800行,后端约1700行)"),
        ("主要用途", "通过数字化手段实现康普茶生产全流程的信息化管理，包括原料管理、发酵过程监控、批次追踪、库存管理等功能，提高生产效率、保证产品质量、降低生产成本\n应用行业：食品加工、饮料制造、发酵工程、生产管理")
    ]

    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=2)

    # 填充表格
    for i, (label, content) in enumerate(table_data):
        row = table.rows[i]

        # 设置第一列
        cell1 = row.cells[0]
        cell1.text = label
        set_cell_font(cell1, 8, True)

        # 设置第二列
        cell2 = row.cells[1]
        cell2.text = content
        set_cell_font(cell2, 8)

    # 设置表格样式
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)

    # 保存文档
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"\n[SUCCESS] 申请表已生成!")
    print(f"[FILE] 保存路径: {output_path}")
    print(f"[INFO] 表格行数: {len(table_data)}行")
    print(f"[INFO] 所有信息已填写完成")
    print("="*60)

if __name__ == '__main__':
    generate_application_form_docx()
