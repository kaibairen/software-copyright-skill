#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成传统康普茶数字化生产管理平台V1.0代码手册
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def extract_code_blocks(markdown_content):
    """从Markdown文件中提取所有代码块"""
    code_blocks = []
    in_code_block = False
    current_block = []

    for line in markdown_content.split('\n'):
        if line.strip().startswith('```'):
            if in_code_block:
                code_blocks.append('\n'.join(current_block))
                current_block = []
                in_code_block = False
            else:
                in_code_block = True
        elif in_code_block:
            current_block.append(line)

    return code_blocks

def setup_page_style(section):
    """设置页面样式"""
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.71)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

def setup_header(section, software_name, version):
    """设置页眉"""
    header = section.header
    for para in header.paragraphs:
        para.clear()

    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = 1

    run = para.add_run(f'{software_name} {version}')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')

    pBdr.append(bottom)
    pPr.append(pBdr)

def setup_footer(section):
    """设置页脚（页码）"""
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = 1

    run = para.add_run('页码')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_code_lines(doc, code_lines):
    """添加代码行（删除空行）"""
    for line in code_lines:
        if not line.strip():
            continue

        code_para = doc.add_paragraph()
        code_run = code_para.add_run(line)
        code_run.font.name = '等线'
        code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        code_run.font.size = Pt(10)
        code_para.paragraph_format.line_spacing = 1.04
        code_para.paragraph_format.space_after = Pt(0)
        code_para.paragraph_format.space_before = Pt(0)

def generate_code_manual():
    """生成代码手册主函数"""
    print("="*60)
    print("传统康普茶数字化生产管理平台V1.0 - 代码手册生成器")
    print("="*60)

    software_name = '传统康普茶数字化生产管理平台'
    version = 'V1.0'
    frontend_path = 'code/frontend-code.md'
    backend_path = 'code/backend-code.md'
    output_path = 'materials/传统康普茶数字化生产管理平台V1.0代码手册.docx'

    print("\n[1/4] 读取代码文件...")
    with open(frontend_path, 'r', encoding='utf-8') as f:
        frontend_content = f.read()

    with open(backend_path, 'r', encoding='utf-8') as f:
        backend_content = f.read()

    print("[2/4] 提取代码块...")
    frontend_blocks = extract_code_blocks(frontend_content)
    backend_blocks = extract_code_blocks(backend_content)

    print(f"  - 前端代码块: {len(frontend_blocks)}个")
    print(f"  - 后端代码块: {len(backend_blocks)}个")

    print("[3/4] 整理代码...")
    all_non_empty_lines = []

    for i, block in enumerate(frontend_blocks):
        lines = block.split('\n')
        non_empty = [line for line in lines if line.strip()]
        all_non_empty_lines.extend(non_empty)

    for i, block in enumerate(backend_blocks):
        lines = block.split('\n')
        non_empty = [line for line in lines if line.strip()]
        all_non_empty_lines.extend(non_empty)

    total_lines = len(all_non_empty_lines)
    print(f"  - 总代码行数: {total_lines}行")

    max_lines = 3000
    selected_lines = all_non_empty_lines[:max_lines]

    print("[4/4] 生成Word文档...")
    doc = Document()

    setup_page_style(doc.sections[0])
    setup_header(doc.sections[0], software_name, version)
    setup_footer(doc.sections[0])

    add_code_lines(doc, selected_lines)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print("\n" + "="*60)
    print("[SUCCESS] 代码手册生成完成!")
    print("="*60)
    print(f"[STATS] 统计信息:")
    print(f"   - 代码行数: {len(selected_lines)}行")
    print(f"   - 目标页数: 60页（每页50行）")
    print(f"   - 字体: 等线 10号")
    print(f"   - 行距: 1.04倍")
    print(f"   - 空行: 已全部删除")
    print(f"\n[FILE] 保存路径:")
    print(f"   {output_path}")
    print("="*60)

if __name__ == '__main__':
    generate_code_manual()
