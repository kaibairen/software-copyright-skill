#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
填写软著申请表
保留模板中已有内容，只填写空白条目
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os


def set_cell_font(cell):
    """
    设置单元格内容的字体为宋体8号

    Args:
        cell: 表格单元格对象
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(8)


def read_application_template(template_path):
    """
    读取申请表模板docx文件的结构
    返回表格的行数和列数，以及每个单元格的标签

    Args:
        template_path: 模板文件路径

    Returns:
        dict: 包含表格结构信息的字典
    """
    doc = Document(template_path)

    if len(doc.tables) == 0:
        print("模板中没有找到表格")
        return None

    table = doc.tables[0]
    print(f"表格行数: {len(table.rows)}, 列数: {len(table.columns)}")

    # 读取第一列的标签
    labels = {}
    for i, row in enumerate(table.rows):
        if len(row.cells) >= 1:
            label = row.cells[0].text.strip()
            if label:
                labels[i] = label
                print(f"行{i}标签: {label}")

    return {
        'rows': len(table.rows),
        'columns': len(table.columns),
        'labels': labels
    }


def generate_application_form(template_path, project_info, output_path):
    """
    根据模板填写申请表（保留已有内容，只填写空白条目）

    Args:
        template_path: 申请表模板路径
        project_info: 项目信息字典，包含以下键：
            - software_name: 软件全称
            - version: 版本号
            - short_name: 软件简称
            - software_type: 软件类型
            - completion_date: 首次发表日期
            - dev_hardware: 开发硬件环境
            - runtime_hardware: 运行硬件环境
            - dev_os: 开发操作系统
            - dev_tools: 开发工具/编程语言
            - runtime_platform: 软件运行平台/操作系统
            - runtime_support: 软件运行支撑环境/支撑软件
            - programming_language: 编程语言
            - code_lines: 源程序量
            - dev_purpose: 开发目的
            - application_field: 应用行业/领域
            - main_functions: 软件的主要功能
            - technical_features: 软件的技术特点
        output_path: 输出文件路径

    Returns:
        str: 输出文件路径
    """
    # 读取模板
    doc = Document(template_path)
    table = doc.tables[0]

    # 首先设置整个表格的字体为宋体8号
    for row in table.rows:
        for cell in row.cells:
            set_cell_font(cell)

    # 定义填写数据映射（行索引 -> 内容）
    fill_data = {
        1: project_info.get('software_name', ''),              # 软件全称
        2: project_info.get('version', ''),                     # 版本号
        3: project_info.get('short_name', ''),                  # 软件简称
        4: project_info.get('software_type', ''),               # 软件类型
        5: project_info.get('completion_date', ''),             # 首次发表日期
        11: project_info.get('dev_hardware', ''),               # 开发硬件环境
        12: project_info.get('runtime_hardware', ''),           # 运行硬件环境
        13: project_info.get('dev_os', ''),                     # 开发操作系统
        14: project_info.get('dev_tools', ''),                  # 开发工具/编程语言
        15: project_info.get('runtime_platform', ''),           # 软件运行平台/操作系统
        16: project_info.get('runtime_support', ''),            # 软件运行支撑环境/支撑软件
        17: project_info.get('programming_language', ''),        # 编程语言
        18: project_info.get('code_lines', ''),                 # 源程序量
        19: project_info.get('dev_purpose', ''),                # 开发目的
        20: project_info.get('application_field', ''),           # 应用行业/领域
        21: project_info.get('main_functions', ''),              # 软件的主要功能
        22: project_info.get('technical_features', ''),          # 软件的技术特点
    }

    # 填充表格（只填写空白条目，保留已有内容）
    filled_count = 0
    for row_idx, text in fill_data.items():
        if row_idx < len(table.rows):
            cell = table.rows[row_idx].cells[1]
            current_value = cell.text.strip()

            # 只在空白时填写
            if not current_value or current_value == '':
                cell.text = str(text)
                # 设置填写内容的字体为宋体8号
                set_cell_font(cell)
                filled_count += 1
                print(f"[NEW] 已填写行{row_idx}")
            else:
                print(f"[KEEP] 保留行{row_idx} (已有内容)")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 保存文件
    doc.save(output_path)
    print(f"申请表已保存至: {output_path}")
    print(f"新填写条目: {filled_count}个")
    return output_path


if __name__ == '__main__':
    # 使用说明：
    # 1. 确保申请表模板文件存在（申请表模板.docx）
    # 2. 准备项目信息字典，包含所有必需的字段
    # 3. 根据实际项目修改 project_info 参数
    #
    # 示例：
    # project_info = {
    #     'software_name': '您的软件全称',
    #     'version': 'V1.0',
    #     'short_name': '软件简称',
    #     'software_type': '应用软件',
    #     'completion_date': '2025年01月21日',
    #     'dev_hardware': '四核intel i5 CPU，16GB内存，512固态硬盘，10Mbps网络带宽',
    #     'runtime_hardware': '服务端内存2G以上，硬盘空间不低于40G； 客户端4G及以上内存，硬盘空间16G及以上',
    #     'dev_os': 'Windows 10',
    #     'dev_tools': 'Python 3.8+, Flask 2.3+',
    #     'runtime_platform': 'Windows 10',
    #     'runtime_support': 'Python运行环境, SQLite/MySQL数据库',
    #     'programming_language': 'Python(后端核心), HTML5(前端)',
    #     'code_lines': '约3000行(前端约1500行,后端约1500行)',
    #     'dev_purpose': '开发目的描述',
    #     'application_field': '应用行业/领域',
    #     'main_functions': '主要功能描述（用分号分隔）',
    #     'technical_features': '技术特点描述'
    # }
    #
    # generate_application_form(
    #     template_path='../asserts/申请表模板.docx',
    #     project_info=project_info,
    #     output_path='./materials/申请表.docx'
    # )
    pass
